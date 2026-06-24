#!/usr/bin/env python3
"""Generate an EDITABLE .docx revision proposal for Dr. Buckholtz (Russian draft).

Structure: cover note + 3-column comparison table (Было / Предлагаем / Почему,
where 'Почему' draws on the reviewer HATE/LOVE pattern) + ranked open questions
+ a 'what we verified' block. Reproducible: re-run to regenerate the .docx.

Output: paper/revision_proposal_for_TJB_RU.docx
"""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor


def shade_cell(cell, hex_fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.makeelement(qn("w:shd"), {qn("w:fill"): hex_fill})
    tc_pr.append(shd)


def h(doc, text, size=14, color=(0x1F, 0x4E, 0x79)):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(size)
    r.font.color.rgb = RGBColor(*color)
    return p


doc = Document()
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(11)

# ---- Title ----
t = doc.add_paragraph()
t.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = t.add_run("Предложение по доработке препринта — черновик на согласование")
r.bold = True
r.font.size = Pt(16)
sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub.add_run("Для д-ра Томаса Дж. Бакхольца · препринт v6 (preprints202511.0598.v6)").italic = True

# ---- Cover note ----
h(doc, "Сопроводительная записка")
cover = (
    "Уважаемый д-р Бакхольц,\n\n"
    "Я взял на себя инициативу: изучил, как структурированы сопоставимые препринты "
    "по физике элементарных частиц, независимо перепроверил числовые соотношения "
    "вашего препринта v6 (каждое число воспроизводится запускаемым скриптом) и "
    "подготовил предложение по доработке.\n\n"
    "Это ЧЕРНОВИК для совместной работы, а не готовый вердикт — вы можете править "
    "его прямо в этом файле.\n\n"
    "Что внутри: (1) сравнительная таблица «было → предлагаем → почему», где «почему» "
    "опирается на то, за что рецензенты обычно заворачивают работу; (2) открытые "
    "вопросы, отмеченные явно, чтобы решить их вместе; (3) краткий список того, что я "
    "уже проверил и на чём предложение стоит.\n\n"
    "Ключевая рекомендация: разделить материал на (А) короткую строгую заметку о "
    "проверенных массовых соотношениях — она ничего не оверклеймит и имеет реальный "
    "шанс пройти рецензию; и (Б) программную статью о космологии MULTING, где открытые "
    "вопросы честно обозначены как направление работы.\n\n"
    "С уважением,\nСергей Бойко · Ronin Institute"
)
for para in cover.split("\n\n"):
    doc.add_paragraph(para)

# ---- Comparison table ----
h(doc, "1. Сравнительная таблица изменений")
rows = [
    ("Было (v6)", "Предлагаем", "Почему"),
    (
        "9+ таблиц; уравнения внутри ячеек таблиц.",
        "Уравнения вынесены в отдельные нумерованные формулы; таблицы — только для данных.",
        "Рецензент должен видеть формулу сразу, а не искать её в ячейке. «Скрытые» уравнения читаются как недостаток строгости.",
    ),
    (
        "Структура нестандартная; раздел «Authority that our work suggests».",
        "Стандартная структура: Аннотация → Введение → Формализм → Результаты → Обсуждение (ограничения) → Выводы.",
        "Рецензенты HATE: нестандартная подача и формулировки, звучащие как заявка на «авторитет». LOVE: знакомый каркас, по которому легко судить работу.",
    ),
    (
        "Термины MESI / IDM / OM / MULTING без определений при первом упоминании.",
        "Каждый термин определяется при первом упоминании; жаргон минимизирован.",
        "Читаемость — первое, на чём заворачивают. Незнакомый жаргон без определений создаёт барьер уже на 1-й странице.",
    ),
    (
        "Проверяемые соотношения и спекулятивная космология идут вместе.",
        "Раздел А — Проверенное (Eq.32, 7:9:17, спектр фермионов). Раздел Б — Программа (космология, β).",
        "Рецензенты LOVE честное отделение крепкого от спекулятивного. Сильный результат не должен тонуть рядом с незакрытыми гипотезами.",
    ),
    (
        "β_d, β_q вводятся как данность (формулы 18–20).",
        "β вынесены в явный блок «Открытый вопрос» с нашим Fisher-анализом неидентифицируемости.",
        "Рецензенты HATE скрытые/подгоночные параметры. Честное «это открытый вопрос + вот анализ» сильнее, чем подача параметра как известного.",
    ),
    (
        "Числа без воспроизводимого источника.",
        "Каждое число — со ссылкой на запускаемый скрипт (verify_all_claims.py и др.); код в открытом репозитории.",
        "Рецензенты LOVE воспроизводимость (GitHub). Это снимает половину придирок «откуда число?» до того, как они заданы.",
    ),
]
table = doc.add_table(rows=len(rows), cols=3)
table.style = "Table Grid"
table.alignment = WD_TABLE_ALIGNMENT.CENTER
for j, txt in enumerate(rows[0]):
    c = table.rows[0].cells[j]
    c.text = ""
    rr = c.paragraphs[0].add_run(txt)
    rr.bold = True
    rr.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    shade_cell(c, "1F4E79")
for i, row in enumerate(rows[1:], start=1):
    for j, txt in enumerate(row):
        table.rows[i].cells[j].text = txt

# ---- Open questions ----
h(doc, "2. Открытые вопросы — решаем вместе")
doc.add_paragraph(
    "Это не замалчивается, а выносится явно. По уровню серьёзности для рецензента "
    "(🔴 CRITICAL — без ответа возможен отказ):"
)
oq = [
    (
        "🔴 Лагранжиан / действие для силового закона F_oP",
        "Сейчас закон постулируется по аналогии с электромагнетизмом. Рецензент спросит: "
        "«покажите вывод из действия или ссылку». Предложение: честно обозначить как открытый "
        "вопрос и отнести в раздел Б (программа).",
    ),
    (
        "🔴 Вывод β_d, β_q из первых принципов",
        "Сейчас 2 свободных параметра. Наш Fisher-анализ: они не идентифицируются из Table A1. "
        "Предложение: пометить как open question; показать анализ — это честнее и сильнее.",
    ),
    (
        "🔴 Аналитический мост F_oP → H(z)",
        "Явного вывода нет (Приложение A.1 — процедура, не вывод). Предложение: обозначить H(z) "
        "как феноменологическое следствие, а не предсказание теории.",
    ),
]
for title, body in oq:
    p = doc.add_paragraph()
    p.add_run(title).bold = True
    doc.add_paragraph(body)

# ---- What we verified ----
h(doc, "3. Что уже проверено (на чём стоит предложение)")
ver = [
    "Eq.32 (4/3)(m_τ/m_e)¹² = α_EM/α_G: отклонение 0.0135% (0.17σ); степень 12 уникальна "
    "(1 из 5040 простых комбинаций); соотношение не найдено в литературе. Скрипт: verify_all_claims.py.",
    "7:9:17 (m_W²:m_Z²:m_H²): держится <0.05%; в scale-free форме m_W/m_H=√(7/17) не зависит от "
    "якоря; предсказывает тяжёлый W (совместимо только с CDF-II/CMS, χ²=2.0; исключает лёгкий, "
    "χ²=39) — фальсифицируемый дискриминатор W-аномалии. Скрипт: c6_mass_preference.py.",
    "Спектр фермионов (Eq.21–24): мюон 0.47%, геометрические средние кварков <0.31% (PDG 2024). "
    "Скрипт: idm_masses.py.",
    "N_opt = ω_cdm/ω_b = 5.366: отклонение от целого 5 равно 5.67σ (исправлена ошибка в формуле "
    "погрешности: было ~259σ из-за несоответствия единиц).",
]
for v in ver:
    doc.add_paragraph(v, style="List Bullet")

doc.add_paragraph()
foot = doc.add_paragraph()
foot.add_run(
    "Черновик подготовлен на основе протоколов рецензентской защиты и pre-submission "
    "QA (Ronin Institute). Английская версия — по запросу."
).italic = True

out = Path(__file__).resolve().parent.parent / "paper" / "revision_proposal_for_TJB_RU.docx"
doc.save(str(out))
print(f"Saved: {out}")
print(
    f"Sections: cover + comparison table ({len(rows) - 1} rows) + {len(oq)} open questions + {len(ver)} verified items"
)
